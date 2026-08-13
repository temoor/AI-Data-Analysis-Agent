import streamlit as st
import pandas as pd
import re
from io import BytesIO

from pypdf import PdfReader
from docx import Document


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="Content Validity Review",
    page_icon="📝",
    layout="wide"
)


# ============================================================
# TITLE
# ============================================================

st.title("📝 Questionnaire Content Validity Review")

st.write(
    "Upload expert evaluation reports and organize the "
    "qualitative review of questionnaire items."
)

st.info(
    "This module records qualitative expert feedback. "
    "It does not calculate CVI, automatically judge validity, "
    "generate new questions, or rewrite questionnaire items."
)


# ============================================================
# HELPER: EXTRACT TEXT FROM PDF
# ============================================================

def extract_pdf_text(uploaded_file):

    uploaded_file.seek(0)

    pdf_bytes = uploaded_file.read()

    reader = PdfReader(
        BytesIO(pdf_bytes)
    )

    pages_text = []

    for page in reader.pages:

        try:
            page_text = page.extract_text()

            if page_text:
                pages_text.append(page_text)

        except Exception:
            continue

    return "\n".join(pages_text)


# ============================================================
# HELPER: EXTRACT TEXT FROM WORD
# ============================================================

def extract_docx_text(uploaded_file):

    uploaded_file.seek(0)

    document = Document(
        uploaded_file
    )

    all_text = []

    # Normal paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            all_text.append(text)

    # Tables
    for table in document.tables:

        for row in table.rows:

            cells = []

            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    cells.append(text)

            if cells:
                all_text.append(
                    " | ".join(cells)
                )

    return "\n".join(all_text)


# ============================================================
# HELPER: CLEAN TEXT
# ============================================================

def clean_text(text):

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


# ============================================================
# QUESTION CODE DETECTION
# ============================================================

def detect_question_code(text):

    """
    Detect common questionnaire codes.

    Examples:
    EP1
    EP2
    SD1
    SD2
    C1
    C2-1
    C2_1
    Q1
    Q2
    A1
    """

    patterns = [

        # Example: C2-1, C2_1
        r"\b[A-Za-z]{1,10}\d{1,3}[-_]\d{1,3}\b",

        # Example: EP1, SD1, A1, Q1
        r"\b[A-Za-z]{1,10}\d{1,3}\b",

    ]

    for pattern in patterns:

        match = re.search(
            pattern,
            text
        )

        if match:

            return match.group(0)

    return None


# ============================================================
# EXTRACT POSSIBLE REVIEW RECORDS
# ============================================================

def extract_review_records(text):

    """
    Flexible extraction of questionnaire review information.

    The function attempts to identify:
    - Question code
    - Questionnaire statement
    - Expert comment

    Because expert reports can use different wording and layouts,
    the original text is preserved where automatic separation
    is uncertain.
    """

    lines = text.splitlines()

    records = []

    current_code = None
    current_statement = []
    current_comment = []

    inside_comment = False
    inside_statement = False

    # --------------------------------------------------------
    # Save current record
    # --------------------------------------------------------

    def save_record():

        nonlocal current_code
        nonlocal current_statement
        nonlocal current_comment

        if not current_code:
            return

        statement = clean_text(
            " ".join(current_statement)
        )

        comment = clean_text(
            " ".join(current_comment)
        )

        # If there is no clear statement, use available text
        if not statement:

            statement = ""

        if not comment:

            comment = ""

        records.append(
            {
                "Question Code": current_code,
                "Question / Statement": statement,
                "Expert Comment": comment
            }
        )

        current_code = None
        current_statement = []
        current_comment = []

    # --------------------------------------------------------
    # Process lines
    # --------------------------------------------------------

    for raw_line in lines:

        line = raw_line.strip()

        if not line:
            continue

        # Remove excessive table separators
        line = line.replace(
            "\t",
            " "
        )

        # ----------------------------------------------------
        # COMMENT LABEL
        # ----------------------------------------------------

        comment_match = re.match(
            r"^(expert\s*)?"
            r"(comment|comments|feedback|remarks|"
            r"recommendation|recommendations|"
            r"observation|observations|"
            r"decision|reviewer\s*comment)"
            r"\s*[:\-]?\s*(.*)$",
            line,
            flags=re.IGNORECASE
        )

        if comment_match:

            inside_comment = True
            inside_statement = False

            comment_text = comment_match.group(3).strip()

            if comment_text:
                current_comment.append(
                    comment_text
                )

            continue

        # ----------------------------------------------------
        # QUESTION / STATEMENT LABEL
        # ----------------------------------------------------

        statement_match = re.match(
            r"^(question|statement|item|questionnaire\s*item)"
            r"\s*[:\-]?\s*(.*)$",
            line,
            flags=re.IGNORECASE
        )

        if statement_match:

            inside_statement = True
            inside_comment = False

            statement_text = (
                statement_match.group(2).strip()
            )

            if statement_text:
                current_statement.append(
                    statement_text
                )

            continue

        # ----------------------------------------------------
        # QUESTION CODE
        # ----------------------------------------------------

        code = detect_question_code(
            line
        )

        # Determine whether the line begins with a code
        code_start_match = re.match(
            r"^\s*"
            r"([A-Za-z]{1,10}\d{1,3}"
            r"(?:[-_]\d{1,3})?)"
            r"\s*[\:\.\-\)]?\s*(.*)$",
            line
        )

        if code_start_match:

            new_code = (
                code_start_match.group(1)
            )

            remaining = (
                code_start_match.group(2).strip()
            )

            # Save previous record
            if current_code:

                save_record()

            current_code = new_code

            inside_statement = True
            inside_comment = False

            if remaining:

                current_statement.append(
                    remaining
                )

            continue

        # ----------------------------------------------------
        # NUMBERED ITEMS
        # ----------------------------------------------------

        numbered_match = re.match(
            r"^\s*(\d{1,3})"
            r"\s*[\.\)\-:]\s*(.+)$",
            line
        )

        if numbered_match:

            if current_code:

                save_record()

            current_code = (
                "Q"
                + numbered_match.group(1)
            )

            current_statement.append(
                numbered_match.group(2).strip()
            )

            inside_statement = True
            inside_comment = False

            continue

        # ----------------------------------------------------
        # CONTINUE COMMENT
        # ----------------------------------------------------

        if inside_comment and current_code:

            current_comment.append(
                line
            )

            continue

        # ----------------------------------------------------
        # CONTINUE STATEMENT
        # ----------------------------------------------------

        if inside_statement and current_code:

            current_statement.append(
                line
            )

    # Save final record
    if current_code:

        save_record()

    # --------------------------------------------------------
    # Remove duplicate records
    # --------------------------------------------------------

    cleaned_records = []

    seen = set()

    for record in records:

        code = record[
            "Question Code"
        ]

        statement = record[
            "Question / Statement"
        ]

        comment = record[
            "Expert Comment"
        ]

        key = (
            code,
            statement,
            comment
        )

        if key not in seen:

            seen.add(key)

            cleaned_records.append(
                record
            )

    return cleaned_records


# ============================================================
# SESSION STATE
# ============================================================

if "expert_reports" not in st.session_state:

    st.session_state.expert_reports = []


if "expert_reviews" not in st.session_state:

    st.session_state.expert_reviews = []


# ============================================================
# UPLOAD EXPERT REPORTS
# ============================================================

st.subheader(
    "📂 Upload Expert Evaluation Reports"
)

st.write(
    "Upload the completed evaluation reports provided by "
    "your academic and industrial experts."
)

uploaded_files = st.file_uploader(
    "Upload one or more expert reports",
    type=["pdf", "docx"],
    accept_multiple_files=True,
    help=(
        "You can upload multiple PDF or Word expert "
        "evaluation reports at the same time."
    )
)


# ============================================================
# PROCESS UPLOADED REPORTS
# ============================================================

if uploaded_files:

    st.success(
        f"✅ {len(uploaded_files)} expert report(s) uploaded."
    )

    st.subheader(
        "👥 Identify Expert Reports"
    )

    expert_information = []

    for index, uploaded_file in enumerate(
        uploaded_files
    ):

        st.markdown(
            f"### Expert Report {index + 1}"
        )

        col1, col2, col3 = st.columns(
            [2, 2, 3]
        )

        with col1:

            expert_name = st.text_input(
                "Expert",
                value=f"Expert {index + 1}",
                key=f"expert_name_{index}"
            )

        with col2:

            expert_type = st.selectbox(
                "Expert Type",
                [
                    "Academic",
                    "Industrial"
                ],
                key=f"expert_type_{index}"
            )

        with col3:

            st.write(
                "📄 Uploaded File"
            )

            st.code(
                uploaded_file.name
            )

        expert_information.append(
            {
                "Expert": expert_name,
                "Expert Type": expert_type,
                "File Name": uploaded_file.name
            }
        )

    st.divider()

    # --------------------------------------------------------
    # PROCESS BUTTON
    # --------------------------------------------------------

    if st.button(
        "🔍 Extract Expert Evaluations",
        type="primary"
    ):

        all_extracted = []

        progress = st.progress(
            0
        )

        for index, uploaded_file in enumerate(
            uploaded_files
        ):

            try:

                filename = uploaded_file.name.lower()

                if filename.endswith(".pdf"):

                    text = extract_pdf_text(
                        uploaded_file
                    )

                elif filename.endswith(".docx"):

                    text = extract_docx_text(
                        uploaded_file
                    )

                else:

                    text = ""

                records = extract_review_records(
                    text
                )

                expert = expert_information[
                    index
                ]

                for record in records:

                    record["Expert"] = (
                        expert["Expert"]
                    )

                    record["Expert Type"] = (
                        expert["Expert Type"]
                    )

                    record["Source File"] = (
                        expert["File Name"]
                    )

                    all_extracted.append(
                        record
                    )

            except Exception as error:

                st.error(
                    f"Could not process "
                    f"{uploaded_file.name}: "
                    f"{error}"
                )

            progress.progress(
                int(
                    ((index + 1)
                    / len(uploaded_files))
                    * 100
                )
            )

        st.session_state[
            "expert_reports"
        ] = all_extracted

        st.success(
            f"✅ Extraction completed. "
            f"{len(all_extracted)} possible "
            f"review record(s) detected."
        )


# ============================================================
# SHOW EXTRACTED EXPERT EVALUATIONS
# ============================================================

extracted_reports = st.session_state.get(
    "expert_reports",
    []
)

if extracted_reports:

    st.subheader(
        "🔍 Extracted Expert Evaluations"
    )

    st.caption(
        "The original expert comment is preserved. "
        "Please check the extracted information before "
        "recording your final researcher decision."
    )

    extracted_df = pd.DataFrame(
        extracted_reports
    )

    display_columns = [
        "Expert",
        "Expert Type",
        "Question Code",
        "Question / Statement",
        "Expert Comment"
    ]

    available_columns = [
        column
        for column in display_columns
        if column in extracted_df.columns
    ]

    st.dataframe(
        extracted_df[
            available_columns
        ],
        use_container_width=True,
        hide_index=True
    )


# ============================================================
# RESEARCHER REVIEW
# ============================================================

if extracted_reports:

    st.divider()

    st.subheader(
        "👨‍🏫 Researcher Review"
    )

    st.write(
        "Review the expert's original comment and "
        "then record your own final decision."
    )

    # --------------------------------------------------------
    # Unique item labels
    # --------------------------------------------------------

    item_options = []

    for record in extracted_reports:

        code = record[
            "Question Code"
        ]

        statement = record[
            "Question / Statement"
        ]

        label = (
            f"{code} — {statement}"
            if statement
            else code
        )

        if label not in item_options:

            item_options.append(
                label
            )

    selected_item = st.selectbox(
        "Select Questionnaire Item",
        item_options
    )

    selected_code = (
        selected_item.split(" — ")[0]
    )

    # --------------------------------------------------------
    # Display all expert comments for item
    # --------------------------------------------------------

    matching_records = [

        record
        for record in extracted_reports
        if record[
            "Question Code"
        ] == selected_code

    ]

    if matching_records:

        first_record = (
            matching_records[0]
        )

        st.markdown(
            "### 📝 Questionnaire Statement"
        )

        st.info(
            first_record[
                "Question / Statement"
            ]
            if first_record[
                "Question / Statement"
            ]
            else
            "Statement was not clearly separated "
            "during automatic extraction."
        )

        st.markdown(
            "### 💬 Expert Comments"
        )

        for record in matching_records:

            expert_label = (
                f"{record['Expert']} "
                f"({record['Expert Type']})"
            )

            with st.expander(
                f"👤 {expert_label}"
            ):

                st.write(
                    "**Original Expert Comment:**"
                )

                comment = record[
                    "Expert Comment"
                ]

                if comment:

                    st.write(
                        comment
                    )

                else:

                    st.warning(
                        "No separate expert comment "
                        "was automatically detected."
                    )

                st.caption(
                    f"Source: {record['Source File']}"
                )

        # ----------------------------------------------------
        # Researcher Decision Dropdown
        # ----------------------------------------------------

        st.markdown(
            "### 👨‍🏫 Researcher Decision"
        )

        researcher_decision = st.selectbox(
            "Select your final decision",
            [
                "Pending Review",
                "Retain",
                "Revise",
                "Remove"
            ],
            key=f"decision_{selected_code}"
        )

        researcher_note = st.text_area(
            "Researcher Note (Optional)",
            placeholder=(
                "Briefly explain your decision "
                "if necessary."
            ),
            key=f"note_{selected_code}"
        )

        if st.button(
            "💾 Save Researcher Decision",
            type="primary",
            key=f"save_{selected_code}"
        ):

            # Remove previous decision for same item
            st.session_state[
                "expert_reviews"
            ] = [

                review
                for review in
                st.session_state[
                    "expert_reviews"
                ]
                if review[
                    "Question Code"
                ] != selected_code

            ]

            new_review = {

                "Question Code":
                    selected_code,

                "Question / Statement":
                    first_record[
                        "Question / Statement"
                    ],

                "Expert Count":
                    len(matching_records),

                "Academic Experts":
                    sum(
                        1
                        for record
                        in matching_records
                        if record[
                            "Expert Type"
                        ] == "Academic"
                    ),

                "Industrial Experts":
                    sum(
                        1
                        for record
                        in matching_records
                        if record[
                            "Expert Type"
                        ] == "Industrial"
                    ),

                "Researcher Decision":
                    researcher_decision,

                "Researcher Note":
                    researcher_note
            }

            st.session_state[
                "expert_reviews"
            ].append(
                new_review
            )

            st.success(
                f"✅ Researcher decision saved for "
                f"{selected_code}."
            )


# ============================================================
# REVIEW SUMMARY
# ============================================================

saved_reviews = st.session_state.get(
    "expert_reviews",
    []
)

if saved_reviews:

    st.divider()

    st.subheader(
        "📊 Researcher Decision Summary"
    )

    summary_df = pd.DataFrame(
        saved_reviews
    )

    st.dataframe(
        summary_df,
        use_container_width=True,
        hide_index=True
    )

    # --------------------------------------------------------
    # Decision counts
    # --------------------------------------------------------

    st.subheader(
        "📈 Decision Overview"
    )

    decision_counts = (
        summary_df[
            "Researcher Decision"
        ]
        .value_counts()
    )

    col1, col2, col3, col4 = st.columns(
        4
    )

    with col1:

        st.metric(
            "Retain",
            int(
                decision_counts.get(
                    "Retain",
                    0
                )
            )
        )

    with col2:

        st.metric(
            "Revise",
            int(
                decision_counts.get(
                    "Revise",
                    0
                )
            )
        )

    with col3:

        st.metric(
            "Remove",
            int(
                decision_counts.get(
                    "Remove",
                    0
                )
            )
        )

    with col4:

        st.metric(
            "Pending",
            int(
                decision_counts.get(
                    "Pending Review",
                    0
                )
            )
        )


# ============================================================
# COVERAGE CHECK
# ============================================================

if extracted_reports:

    st.divider()

    st.subheader(
        "🔎 Expert Review Coverage"
    )

    total_experts = len(
        uploaded_files
    )

    coverage_data = []

    unique_codes = sorted(
        set(
            record[
                "Question Code"
            ]
            for record
            in extracted_reports
        )
    )

    for code in unique_codes:

        code_records = [

            record
            for record in extracted_reports
            if record[
                "Question Code"
            ] == code

        ]

        expert_count = len(
            set(
                record[
                    "Expert"
                ]
                for record
                in code_records
            )
        )

        academic_count = len(
            set(
                record[
                    "Expert"
                ]
                for record
                in code_records
                if record[
                    "Expert Type"
                ] == "Academic"
            )
        )

        industrial_count = len(
            set(
                record[
                    "Expert"
                ]
                for record
                in code_records
                if record[
                    "Expert Type"
                ] == "Industrial"
            )
        )

        coverage_data.append(
            {
                "Question Code":
                    code,

                "Experts Reviewed":
                    expert_count,

                "Academic":
                    academic_count,

                "Industrial":
                    industrial_count,

                "Expected Experts":
                    total_experts,

                "Coverage":
                    (
                        "Complete"
                        if expert_count
                        >= total_experts
                        else "Incomplete"
                    )
            }
        )

    if coverage_data:

        coverage_df = pd.DataFrame(
            coverage_data
        )

        st.dataframe(
            coverage_df,
            use_container_width=True,
            hide_index=True
        )


# ============================================================
# METHODOLOGICAL NOTE
# ============================================================

st.divider()

st.subheader(
    "📌 Methodological Note"
)

st.write(
    "This module is designed to document qualitative "
    "expert review of questionnaire items. Expert comments "
    "are preserved in their original form. The researcher "
    "makes the final decision to retain, revise, remove, "
    "or keep an item pending review. The module does not "
    "calculate Content Validity Index (CVI)."
)
